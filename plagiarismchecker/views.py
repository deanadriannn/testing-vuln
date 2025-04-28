from django.shortcuts import render
from django.http import HttpResponse
from plagiarismchecker.algorithm import main
from docx import *
from plagiarismchecker.algorithm import fileSimilarity
import PyPDF2 
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

# Create your views here.
#home
def home(request):
    return render(request, 'pc/index.html') 

#web search(Text)
def test(request):
    print("request is welcome test")
    print(request.POST['q'])  
    
    if request.POST['q']: 
        percent,link = main.findSimilarity(request.POST['q'])
        percent = round(percent,2)
    print("Output.....................!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent})

#web search file(.txt, .docx)
# !ORIGINAL
# def filetest(request):
#     value = ''    
#     print("REQUEST DARI USER: ", request)
#     print(request.FILES['docfile'])
#     if str(request.FILES['docfile']).endswith(".txt"):
#         value = str(request.FILES['docfile'].read())

#     elif str(request.FILES['docfile']).endswith(".docx"):
#         document = Document(request.FILES['docfile'])
#         for para in document.paragraphs:
#             value += para.text

#     elif str(request.FILES['docfile']).endswith(".pdf"):
#         # creating a pdf file object 
#         pdfFileObj = open(request.FILES['docfile'], 'rb') 

#         # creating a pdf reader object 
#         pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 

#         # printing number of pages in pdf file 
#         print(pdfReader.numPages) 

#         # creating a page object 
#         pageObj = pdfReader.getPage(0) 

#         # extracting text from page 
#         print(pageObj.extractText()) 

#         # closing the pdf file object 
#         pdfFileObj.close() 


#     percent,link = main.findSimilarity(value)
#     print("Output...................!!!!!!!!",percent,link)
#     return render(request, 'pc/index.html',{'link': link, 'percent': percent})

@csrf_exempt
def filetest(request):
    if request.method != 'POST':
        return JsonResponse({"error": "Method not allowed"}, status=405)

    if 'docfile' not in request.FILES:
        return JsonResponse({"error": "No file uploaded"}, status=400)

    file = request.FILES['docfile']
    value = ""

    # Proses file sesuai dengan ekstensi
    if file.name.endswith(".txt"):
        value = file.read().decode("utf-8")

    elif file.name.endswith(".docx"):
        document = Document(file)
        for para in document.paragraphs:
            value += para.text + "\n"

    elif file.name.endswith(".pdf"):
        pdfReader = PyPDF2.PdfReader(file)
        for page in pdfReader.pages:
            value += page.extract_text() + "\n"

    else:
        return JsonResponse({"error": "Unsupported file type"}, status=400)

    # Proses pengecekan plagiarisme
    percent, link = main.findSimilarity(value)

    # Return JSON response ke Laravel
    return JsonResponse({
        "percent": percent,
        "link": link
    })

#text compare
def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

#two text compare(Text)
def twofiletest1(request):
    print("Submiited text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])
    result = round(result,2)    
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
    

#two text compare(.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())

    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1,value2)
    
    print("Output..................!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
